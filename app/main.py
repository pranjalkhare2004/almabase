"""Main FastAPI application — Structured Questionnaire Answering Tool."""
import os
import tempfile
from pathlib import Path
from typing import List

from dotenv import load_dotenv
load_dotenv(Path(__file__).resolve().parent.parent / ".env")

from fastapi import FastAPI, Depends, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from docx import Document as DocxDocument

from app.database import get_db, init_db
from app.models import User, Questionnaire, Question, ReferenceDocument
from app.auth import router as auth_router, get_current_user, require_auth, RedirectException
from app.utils import extract_text_from_file, split_questions
from app.rag.retrieval import build_index
from app.rag.orchestrator import answer_question

app = FastAPI(title="Questionnaire Answering Tool")
app.mount("/static", StaticFiles(directory="app/static"), name="static")
templates = Jinja2Templates(directory="app/templates")
app.include_router(auth_router)


@app.exception_handler(RedirectException)
async def redirect_exception_handler(request: Request, exc: RedirectException):
    return RedirectResponse(url=exc.url, status_code=303)


@app.on_event("startup")
def startup():
    init_db()


# --- Homepage ---

@app.get("/", response_class=HTMLResponse)
async def homepage(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    return templates.TemplateResponse("index.html", {"request": request, "user": user})


# --- Dashboard ---

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request, user: User = Depends(require_auth), db: Session = Depends(get_db)):
    questionnaires = db.query(Questionnaire).filter(Questionnaire.user_id == user.id).order_by(Questionnaire.created_at.desc()).all()
    return templates.TemplateResponse("dashboard.html", {
        "request": request, "user": user, "questionnaires": questionnaires,
    })


# --- Upload Questionnaire ---

@app.get("/upload-questionnaire", response_class=HTMLResponse)
async def upload_questionnaire_page(request: Request, user: User = Depends(require_auth)):
    return templates.TemplateResponse("upload_questionnaire.html", {"request": request, "user": user})


@app.post("/upload-questionnaire")
async def upload_questionnaire(
    request: Request,
    file: UploadFile = File(...),
    user: User = Depends(require_auth),
    db: Session = Depends(get_db),
):
    if not file.filename.lower().endswith((".pdf", ".txt")):
        return templates.TemplateResponse("upload_questionnaire.html", {
            "request": request, "user": user, "error": "Only PDF and TXT files are supported.",
        })

    content = await file.read()
    try:
        text = extract_text_from_file(content, file.filename)
    except Exception as e:
        return templates.TemplateResponse("upload_questionnaire.html", {
            "request": request, "user": user, "error": f"Error reading file: {str(e)}",
        })

    questions = split_questions(text)
    if not questions:
        return templates.TemplateResponse("upload_questionnaire.html", {
            "request": request, "user": user, "error": "No questions found in the file.",
        })

    questionnaire = Questionnaire(user_id=user.id, filename=file.filename)
    db.add(questionnaire)
    db.commit()
    db.refresh(questionnaire)

    for i, q_text in enumerate(questions):
        question = Question(questionnaire_id=questionnaire.id, text=q_text, order_index=i)
        db.add(question)
    db.commit()

    return RedirectResponse(url=f"/questionnaire/{questionnaire.id}/upload-references", status_code=303)


# --- Upload Reference Documents ---

@app.get("/questionnaire/{q_id}/upload-references", response_class=HTMLResponse)
async def upload_references_page(
    request: Request, q_id: int,
    user: User = Depends(require_auth), db: Session = Depends(get_db),
):
    questionnaire = db.query(Questionnaire).filter(
        Questionnaire.id == q_id, Questionnaire.user_id == user.id
    ).first()
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found.")

    refs = db.query(ReferenceDocument).filter(ReferenceDocument.questionnaire_id == q_id).all()
    return templates.TemplateResponse("upload_references.html", {
        "request": request, "user": user, "questionnaire": questionnaire, "refs": refs,
    })


@app.post("/questionnaire/{q_id}/upload-references")
async def upload_references(
    request: Request, q_id: int,
    files: List[UploadFile] = File(...),
    user: User = Depends(require_auth),
    db: Session = Depends(get_db),
):
    questionnaire = db.query(Questionnaire).filter(
        Questionnaire.id == q_id, Questionnaire.user_id == user.id
    ).first()
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found.")

    for file in files:
        if not file.filename.lower().endswith((".pdf", ".txt")):
            continue
        content = await file.read()
        try:
            text = extract_text_from_file(content, file.filename)
        except Exception:
            continue

        build_index(questionnaire.id, text, file.filename)

        ref = ReferenceDocument(
            user_id=user.id, questionnaire_id=questionnaire.id, filename=file.filename
        )
        db.add(ref)

    db.commit()
    return RedirectResponse(url=f"/questionnaire/{q_id}/upload-references", status_code=303)


# --- Generate Answers ---

@app.post("/questionnaire/{q_id}/generate")
async def generate_answers(
    request: Request, q_id: int,
    user: User = Depends(require_auth),
    db: Session = Depends(get_db),
):
    questionnaire = db.query(Questionnaire).filter(
        Questionnaire.id == q_id, Questionnaire.user_id == user.id
    ).first()
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found.")

    questions = db.query(Question).filter(
        Question.questionnaire_id == q_id
    ).order_by(Question.order_index).all()

    for question in questions:
        try:
            answer, citation = answer_question(question.text, questionnaire.id)
            question.answer = answer
            question.citation = citation
        except Exception as e:
            question.answer = f"Error generating answer: {str(e)}"
            question.citation = ""

    db.commit()
    return RedirectResponse(url=f"/questionnaire/{q_id}/results", status_code=303)


# --- Results / Edit ---

@app.get("/questionnaire/{q_id}/results", response_class=HTMLResponse)
async def results_page(
    request: Request, q_id: int,
    user: User = Depends(require_auth),
    db: Session = Depends(get_db),
):
    questionnaire = db.query(Questionnaire).filter(
        Questionnaire.id == q_id, Questionnaire.user_id == user.id
    ).first()
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found.")

    questions = db.query(Question).filter(
        Question.questionnaire_id == q_id
    ).order_by(Question.order_index).all()

    return templates.TemplateResponse("results.html", {
        "request": request, "user": user,
        "questionnaire": questionnaire, "questions": questions,
    })


@app.post("/question/{question_id}/edit")
async def edit_answer(
    question_id: int,
    answer: str = Form(...),
    user: User = Depends(require_auth),
    db: Session = Depends(get_db),
):
    question = db.query(Question).filter(Question.id == question_id).first()
    if not question:
        raise HTTPException(status_code=404, detail="Question not found.")

    questionnaire = db.query(Questionnaire).filter(
        Questionnaire.id == question.questionnaire_id,
        Questionnaire.user_id == user.id,
    ).first()
    if not questionnaire:
        raise HTTPException(status_code=403, detail="Not authorized.")

    question.answer = answer
    db.commit()
    return RedirectResponse(url=f"/questionnaire/{questionnaire.id}/results", status_code=303)


# --- Export DOCX ---

@app.get("/questionnaire/{q_id}/export")
async def export_docx(
    q_id: int,
    user: User = Depends(require_auth),
    db: Session = Depends(get_db),
):
    questionnaire = db.query(Questionnaire).filter(
        Questionnaire.id == q_id, Questionnaire.user_id == user.id
    ).first()
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found.")

    questions = db.query(Question).filter(
        Question.questionnaire_id == q_id
    ).order_by(Question.order_index).all()

    doc = DocxDocument()
    doc.add_heading("Questionnaire Answers", level=0)
    doc.add_paragraph(f"Source: {questionnaire.filename}")
    doc.add_paragraph("")

    for i, question in enumerate(questions, 1):
        doc.add_heading(f"Q{i}: {question.text}", level=2)
        doc.add_paragraph(question.answer or "No answer generated.")
        if question.citation:
            citation_para = doc.add_paragraph()
            run = citation_para.add_run(f"Citation: {question.citation}")
            run.italic = True
        doc.add_paragraph("")

    filepath = os.path.join(tempfile.gettempdir(), f"questionnaire_{q_id}_export.docx")
    doc.save(filepath)
    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{questionnaire.filename.rsplit('.', 1)[0]}_answers.docx",
    )
